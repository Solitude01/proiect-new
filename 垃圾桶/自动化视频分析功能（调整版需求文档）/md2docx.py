import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from bs4 import BeautifulSoup
import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading

def process_table(doc, table_element):
    """处理HTML表格转换为Word表格"""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    cols = len(rows[0].find_all(['th', 'td']))
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            table.rows[i].cells[j].text = cell.get_text().strip()

def process_inline_elements(paragraph, element):
    """处理段落中的内联元素"""
    for child in element.children:
        if isinstance(child, str):
            paragraph.add_run(child)
        elif child.name == 'strong' or child.name == 'b':
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name == 'em' or child.name == 'i':
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif child.name == 'a':
            run = paragraph.add_run(child.get_text())
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        else:
            paragraph.add_run(child.get_text())

def convert_md_to_docx(md_file, output_file=None):
    """转换Markdown到Word"""
    # 读取Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为HTML
    html = markdown.markdown(md_content, extensions=['extra', 'codehilite', 'tables'])
    soup = BeautifulSoup(html, 'html.parser')
    
    # 创建Word文档
    doc = Document()
    
    # 处理元素
    for element in soup.children:
        if not element.name:
            continue
            
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)
        elif element.name == 'p':
            p = doc.add_paragraph()
            process_inline_elements(p, element)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'pre':
            p = doc.add_paragraph(element.get_text())
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        elif element.name == 'blockquote':
            p = doc.add_paragraph(element.get_text())
            p.style = 'Quote'
        elif element.name == 'table':
            process_table(doc, element)
    
    # 保存
    if not output_file:
        output_file = os.path.splitext(md_file)[0] + '.docx'
    
    doc.save(output_file)
    return output_file

class MarkdownConverterGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Markdown 转 Word 工具")
        self.root.geometry("700x500")
        self.root.resizable(True, True)
        
        # 设置样式
        style = ttk.Style()
        style.theme_use('clam')
        
        # 文件列表
        self.file_list = []
        
        # 创建界面
        self.create_widgets()
        
    def create_widgets(self):
        # 标题
        title_frame = tk.Frame(self.root, bg='#2c3e50', height=60)
        title_frame.pack(fill='x')
        title_frame.pack_propagate(False)
        
        title_label = tk.Label(
            title_frame, 
            text="📄 Markdown 转 Word 转换器",
            font=('Microsoft YaHei UI', 16, 'bold'),
            bg='#2c3e50',
            fg='white'
        )
        title_label.pack(pady=15)
        
        # 按钮区域
        button_frame = tk.Frame(self.root, bg='#ecf0f1', height=70)
        button_frame.pack(fill='x', pady=10)
        
        btn_style = {
            'font': ('Microsoft YaHei UI', 10),
            'width': 15,
            'height': 2,
            'cursor': 'hand2'
        }
        
        self.select_btn = tk.Button(
            button_frame,
            text="📁 选择文件",
            command=self.select_files,
            bg='#3498db',
            fg='white',
            activebackground='#2980b9',
            **btn_style
        )
        self.select_btn.pack(side='left', padx=10)
        
        self.convert_btn = tk.Button(
            button_frame,
            text="🔄 开始转换",
            command=self.start_conversion,
            bg='#27ae60',
            fg='white',
            activebackground='#229954',
            state='disabled',
            **btn_style
        )
        self.convert_btn.pack(side='left', padx=10)
        
        self.clear_btn = tk.Button(
            button_frame,
            text="🗑️ 清空列表",
            command=self.clear_list,
            bg='#e74c3c',
            fg='white',
            activebackground='#c0392b',
            **btn_style
        )
        self.clear_btn.pack(side='left', padx=10)
        
        # 文件列表区域
        list_frame = tk.Frame(self.root)
        list_frame.pack(fill='both', expand=True, padx=10, pady=10)
        
        list_label = tk.Label(
            list_frame,
            text="待转换文件列表:",
            font=('Microsoft YaHei UI', 10, 'bold'),
            anchor='w'
        )
        list_label.pack(fill='x', pady=(0, 5))
        
        # 创建滚动条
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side='right', fill='y')
        
        self.file_listbox = tk.Listbox(
            list_frame,
            font=('Microsoft YaHei UI', 9),
            selectmode='extended',
            yscrollcommand=scrollbar.set,
            bg='#f8f9fa',
            selectbackground='#3498db'
        )
        self.file_listbox.pack(fill='both', expand=True)
        scrollbar.config(command=self.file_listbox.yview)
        
        # 状态栏
        self.status_frame = tk.Frame(self.root, bg='#34495e', height=40)
        self.status_frame.pack(fill='x', side='bottom')
        self.status_frame.pack_propagate(False)
        
        self.status_label = tk.Label(
            self.status_frame,
            text="就绪",
            font=('Microsoft YaHei UI', 9),
            bg='#34495e',
            fg='white',
            anchor='w'
        )
        self.status_label.pack(fill='both', padx=10, pady=8)
        
    def select_files(self):
        """选择文件"""
        files = filedialog.askopenfilenames(
            title="选择 Markdown 文件",
            filetypes=[("Markdown 文件", "*.md"), ("所有文件", "*.*")]
        )
        
        if files:
            for file in files:
                if file not in self.file_list:
                    self.file_list.append(file)
                    filename = os.path.basename(file)
                    self.file_listbox.insert(tk.END, f"📄 {filename}")
            
            self.convert_btn.config(state='normal')
            self.status_label.config(text=f"已选择 {len(self.file_list)} 个文件")
    
    def clear_list(self):
        """清空列表"""
        self.file_list.clear()
        self.file_listbox.delete(0, tk.END)
        self.convert_btn.config(state='disabled')
        self.status_label.config(text="列表已清空")
    
    def start_conversion(self):
        """开始转换"""
        if not self.file_list:
            messagebox.showwarning("警告", "请先选择要转换的文件！")
            return
        
        # 禁用按钮
        self.select_btn.config(state='disabled')
        self.convert_btn.config(state='disabled')
        self.clear_btn.config(state='disabled')
        
        # 在新线程中执行转换
        thread = threading.Thread(target=self.convert_files)
        thread.daemon = True
        thread.start()
    
    def convert_files(self):
        """转换文件"""
        success_count = 0
        fail_count = 0
        
        for i, md_file in enumerate(self.file_list):
            filename = os.path.basename(md_file)
            self.update_status(f"正在转换 ({i+1}/{len(self.file_list)}): {filename}")
            
            try:
                output_file = convert_md_to_docx(md_file)
                success_count += 1
                self.update_listbox(i, f"✅ {filename}")
            except Exception as e:
                fail_count += 1
                self.update_listbox(i, f"❌ {filename} (失败: {str(e)[:30]}...)")
        
        # 转换完成
        self.root.after(0, self.conversion_complete, success_count, fail_count)
    
    def update_status(self, text):
        """更新状态栏"""
        self.root.after(0, lambda: self.status_label.config(text=text))
    
    def update_listbox(self, index, text):
        """更新列表项"""
        def _update():
            self.file_listbox.delete(index)
            self.file_listbox.insert(index, text)
        self.root.after(0, _update)
    
    def conversion_complete(self, success, fail):
        """转换完成"""
        message = f"转换完成！\n\n成功: {success} 个\n失败: {fail} 个"
        messagebox.showinfo("完成", message)
        
        # 重新启用按钮
        self.select_btn.config(state='normal')
        self.convert_btn.config(state='normal')
        self.clear_btn.config(state='normal')
        
        self.status_label.config(text=f"转换完成 - 成功: {success}, 失败: {fail}")

if __name__ == "__main__":
    root = tk.Tk()
    app = MarkdownConverterGUI(root)
    root.mainloop()